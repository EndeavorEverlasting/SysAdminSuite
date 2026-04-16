#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
generate_labels_v8_1.py — Rock-solid label generator (A/B/H/I), 2-up, non-splitting rows.

Behavior (v8 contract):
  • Sheet: first tab only.
  • OLD HOST  → Column A (index 0 after header row)
  • NEW HOST  → Column B (index 1)
  • NEW SITE  → Column H (index 7)
  • NEW ROOM  → Column I (index 8)
  • Two labels per page (page break after every second).
  • Both checkboxes checked: '☑ LabelTaped    ☑ QA Pass'.
  • Rows included only if either host looks like a hostname (regex), default: [A-Z]{3}\d{3}[A-Z]{3}\d{3}

Safeties:
  • Header row is auto-detected; override with --header-row if needed.
  • Dry-run prints a 10-row preview and counts; no docx written.

Deps: pip install pandas python-docx openpyxl
"""

from __future__ import annotations
import argparse
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

DEFAULT_HOST_REGEX = r"[A-Z]{3}\d{3}[A-Z]{3}\d{3}"  # e.g., WNY075EPT123

def norm(v) -> str:
    if pd.isna(v):
        return ""
    s = str(v).strip()
    return "" if s.lower() in {"nan", "none", "null"} else s

def detect_header_row(raw: pd.DataFrame, scan_rows: int = 40) -> int:
    tokens = {"old pc names", "new pc name", "labeled location", "new room location"}
    for i in range(min(scan_rows, len(raw))):
        row_vals = {str(x).strip().lower() for x in raw.iloc[i].tolist()}
        if row_vals & tokens:
            return i
    return 0

def add_cant_split(table_row) -> None:
    tr = table_row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))

def add_label_block(doc: Document, site: str, room: str, old_host: str, new_host: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    add_cant_split(table.rows[0])
    cell = table.cell(0, 0)

    p = cell.add_paragraph()
    r = p.add_run("REMAIN — DO NOT MOVE (Tear Off)")
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = cell.add_paragraph("—" * 22 + "  Tear here  " + "—" * 22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = cell.add_paragraph()
    r = p.add_run("MOVE — WORKSTATION LABEL")
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in [
        ("NEW SITE: ", site),
        ("NEW ROOM: ", room),
        ("OLD HOST: ", old_host),
        ("NEW HOST: ", new_host),
    ]:
        q = cell.add_paragraph()
        a = q.add_run(label); a.bold = True
        q.add_run(value)

    cell.add_paragraph("")  # spacer
    cell.add_paragraph("☑ LabelTaped    ☑ QA Pass")
    # small spacer after block so two fit nicely on the page
    doc.add_paragraph("")

def build_labels_v8(
    excel_path: Path,
    out_path: Path,
    header_row: int | None = None,
    host_regex: str | None = DEFAULT_HOST_REGEX,
    dry_run: bool = False,
) -> dict:
    if not excel_path.exists():
        raise SystemExit(f"[v8.1] Excel not found: {excel_path}")

    raw = pd.read_excel(excel_path, sheet_name=0, header=None)
    hdr = header_row if header_row is not None else detect_header_row(raw)

    header_values = raw.iloc[hdr].tolist()
    df = raw.iloc[hdr + 1 :].reset_index(drop=True)
    df.columns = header_values

    # Fixed positions after header: A/B/H/I
    colA_old = df.iloc[:, 0]
    colB_new = df.iloc[:, 1]
    colH_site = df.iloc[:, 7]
    colI_room = df.iloc[:, 8]

    pat = re.compile(host_regex) if host_regex else None
    mask = pd.Series(True, index=df.index)
    if pat:
        mask &= colA_old.astype(str).str.contains(pat) | colB_new.astype(str).str.contains(pat)

    data = df.loc[mask].copy()

    preview = pd.DataFrame(
        {
            "A_old": colA_old.loc[data.index].astype(str).head(10),
            "B_new": colB_new.loc[data.index].astype(str).head(10),
            "H_site": colH_site.loc[data.index].astype(str).head(10),
            "I_room": colI_room.loc[data.index].astype(str).head(10),
        }
    )

    summary = {
        "sheet_index": 0,
        "header_row": int(hdr),
        "rows_considered": int(len(df)),
        "labels_to_print": int(len(data)),
        "host_regex": host_regex or "",
        "preview_head": preview.to_dict(orient="records"),
    }

    if dry_run:
        return summary

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    on_page = 0
    for idx in data.index:
        old_host = norm(colA_old.loc[idx])
        new_host = norm(colB_new.loc[idx])
        site = norm(colH_site.loc[idx])
        room = norm(colI_room.loc[idx])

        add_label_block(doc, site, room, old_host, new_host)
        on_page += 1
        if on_page == 2:
            doc.add_page_break()
            on_page = 0

    doc.save(out_path)
    summary["output"] = str(out_path)
    return summary

def main():
    ap = argparse.ArgumentParser(
        description="Generate tear-off workstation labels (v8.1: A/B/H/I, 2-up, non-splitting)."
    )
    ap.add_argument("--excel", required=True, help="Path to the Excel live tracker (.xlsx)")
    ap.add_argument("--out", default="Generated_TearOff_Labels_2up.docx", help="Output .docx filepath")
    ap.add_argument("--header-row", type=int, default=None,
                    help="Override detected header row (0-based). If omitted, the script will try to detect it.")
    ap.add_argument("--host-regex", default=DEFAULT_HOST_REGEX,
                    help="Regex for filtering device rows. Use '' to include all rows.")
    ap.add_argument("--dry-run", action="store_true",
                    help="Preview counts + first 10 rows of A/B/H/I. Does NOT write a DOCX.")
    args = ap.parse_args()

    info = build_labels_v8(
        excel_path=Path(args.excel),
        out_path=Path(args.out),
        header_row=args.header_row,
        host_regex=(args.host_regex if args.host_regex != "" else None),
        dry_run=bool(args.dry_run),
    )

    print(f"[v8.1] header_row={info['header_row']}, labels={info['labels_to_print']}, rows={info['rows_considered']}")
    if args.dry_run:
        print("[v8.1] preview (first 10 rows of A/B/H/I):")
        for row in info["preview_head"]:
            print(row)
    else:
        print(f"[v8.1] wrote: {info['output']}")

if __name__ == "__main__":
    main()
